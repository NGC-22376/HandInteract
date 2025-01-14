"""
从“手语采样数据集.docx”得到数据集目录：单手势词语、多手势词语
新建文档，重排“手语采样数据集.docx”，依照单手势和多手势分类，以便测数据时参照
"""
import os
import re
from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm


def docx_construct_dataset(docx_path, dataset_path):
    """
    构建数据集目录，并新建手势数分类的文档。
    执行完后，数据集格式：单手势/动作名，多手势/动作名，并且在docx_path的同一父目录下新建用手势数分类的文档，其中包含词语名-手势示意图。
    :param docx_path:"手语采样数据集.docx"存在的路径
    :param dataset_path:数据集建立的位置。如：C:/.../数据集
    :return: 无返回值
    """
    # 处理文字，建立词义索引
    doc = Document(docx_path)
    key_list = []
    simple_key_list, complicate_key_list = [], []
    dictionary = {}
    for paragraph in doc.paragraphs:
        # 忽略空段落
        if paragraph.text.strip():
            # 对段落中的每一块连续格式化文本做处理
            for block in paragraph.runs:
                # 提取每张图片对应的词义，建立在词义作为二级标题，字体大小固定为三号字的基础上
                if block.font.size == 203200:
                    key_list.append(block.text)
                    dictionary[block.text] = None

    # 处理图片，加入字典
    for rel in doc.part.rels.values():
        print(rel.target_ref)
        if "image" in rel.target_ref:
            index = int(re.findall(r'\d+', rel.target_ref)[0])
            image_data = BytesIO(rel.target_part.blob)
            dictionary[key_list[index - 1]] = image_data

    # 创建数据集目录结构
    simple_path = os.path.join(dataset_path, "单动作手势")
    complicate_path = os.path.join(dataset_path, "多动作手势")
    os.makedirs(simple_path, exist_ok=True)
    os.makedirs(complicate_path, exist_ok=True)
    for word in key_list:
        if '2' == word[-1]:
            name = word[:-1]
            obj_path = os.path.join(complicate_path, name)  # 匹配数字之前的字符串内容
            complicate_key_list.append(word)
        else:
            obj_path = os.path.join(simple_path, word)
            simple_key_list.append(word)
        if os.path.exists(obj_path):
            print(f"'{word}'重复记录")
        else:
            os.makedirs(obj_path)
            print(f"'{word}'成功建立")

    # 新建以手势个数分类的docx
    new_doc = Document()
    # 设置样式
    styles = new_doc.styles
    type_style = styles.add_style("TypeStyle", WD_STYLE_TYPE.PARAGRAPH)  # 大类的字体样式
    type_style.font.name = "Arial"  # 字体
    type_style.font.size = 228600  # 二号字体
    type_style.font.bold = True  # 加粗
    key_style = styles.add_style("KeyStyle", WD_STYLE_TYPE.PARAGRAPH)  # 图片对应词义的字体样式
    key_style.font.name = "Arial"  # 字体
    key_style.font.size = 203200  # 三号字体
    key_style.font.bold = True  # 加粗
    # 写入单手势动作
    new_doc.add_paragraph(text="单手势动作", style="TypeStyle")
    for key in simple_key_list:
        if dictionary[key] is None:
            continue
        new_doc.add_paragraph(key, style="KeyStyle")
        new_doc.add_picture(dictionary[key], width=Cm(14))
        print(f"'{key}'再分类写入成功")
    # 写入多手势动作
    new_doc.add_paragraph(text="多手势动作", style="TypeStyle")
    for key in complicate_key_list:
        if dictionary[key] is None:
            continue
        new_doc.add_paragraph(key, style="KeyStyle")
        new_doc.add_picture(dictionary[key], width=Cm(14))
        print(f"'{key}'再分类写入成功")

    # 保存文档
    new_doc.save(os.path.join(os.path.dirname(docx_path), "手语采样数据集-按手势数分类.docx"))

# def check_dislocation(docx_path):
#     """
#     检查docx内所有文字和图片是否匹配
#     :param docx_path: docx文件所在路径
#     :return: null
#     """
#     doc = Document(docx_path)
#     key_list = []
#     for paragraph in doc.paragraphs:
#         # 忽略空段落
#         if paragraph.text.strip():
#             # 对段落中的每一块连续格式化文本做处理
#             for block in paragraph.runs:
#                 # 提取每张图片对应的词义，建立在词义作为二级标题，字体大小固定为三号字的基础上
#                 if block.font.size == 203200:
#                     key_list.append(re.match(r"\D+", block.text).group())
#
#     index = 0
#     parts = [doc.part]
#     parts.extend(doc.part.package.parts)
#     for part in parts:
#         for rel in part.rels.values():
#             if "image" in rel.target_ref:
#                 image_data = BytesIO(rel.target_part.blob)
#                 check_match(key_list[index], image_data)
#                 index += 1


# def check_match(text, img):
#     """
#     检查文字与图片是否一一对应
#     :param text: 文字
#     :param img: 字节流格式图片
#     :return: null。输出匹配信息
#     """
#     ocr = PaddleOCR(use_angle_cls=True, lang='ch')  # 支持角度分类和中文
#     result = ocr.ocr(np.array(Image.open(img)), cls=True)  # 传递 Image 对象进行 OCR 识别
#     # 比对识别结果
#     for line in result[0]:
#         if text in line:
#             print(f"{text}匹配成功")
#             return
#     print(f"{text}匹配失败")
#     exit(-1)
