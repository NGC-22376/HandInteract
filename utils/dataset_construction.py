"""
从“手语采样数据集.docx”得到数据集目录：单手势词语、多手势词语
新建文档，重排“手语采样数据集.docx”，依照单手势和多手势分类，以便测数据时参照
数据集格式规整化
"""
import os
import re
from io import BytesIO

import fitz
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm


def extract_csv_data(original_df, new_file, start_row, end_row):
    """
    将csv文件的指定行复制到新文件中
    :param original_df:原始csv数据
    :param new_file: 新数据保存的地址
    :param start_row:转存的起始行数
    :param end_row:转存的结束行数
    :return:null
    """
    if not os.path.exists(os.path.dirname(new_file)):
        os.makedirs(os.path.dirname(new_file))
    # 截取数据并保存为新的 CSV 文件
    new_df = original_df.iloc[start_row:end_row]
    new_df.to_csv(new_file, index=False, encoding='utf-8')  # 不保存索引

    print(f"数据已成功提取到新文件: {new_file}")


def format_folders(dataset_path, data_path):
    """
    实验数据集格式规整化，将乱命名的文件规整为...//手势名//记录xx（间隔时间）。这一步完成后，数据即可进入预处理程序
    :param dataset_path: 数据集父目录。如：C:/.../数据集
    :param data_path: 实验数据父目录。如：E:/.../测试
    """
    for dir_name in os.listdir(data_path):
        index = 0
        dir_path = os.path.join(data_path, dir_name)  # E://...//测试//食指弯曲
        for file_name in os.listdir(dir_path):
            index += 1
            # 匹配间隔时间
            pattern = r'（(.*?)）'
            interval = re.findall(pattern, file_name)[0]
            origin_path = os.path.join(dir_path, file_name)  # 原始文件：E://...//测试//食指弯曲//xxx（1.5）

            if origin_path.endswith('.csv'):
                # 读取 CSV 文件
                original_df = pd.read_csv(origin_path)
                # 获取总行数
                max_time = original_df.iloc[-1, 0]
                total_part = int(max_time // float(interval))
                row_per_part = len(original_df) // total_part
                # 根据间隔时间分割原始文件，并移动到新路径
                for i in range(2, total_part):
                    new_name = f"记录{index}.{i}.csv"
                    new_path = os.path.join(dataset_path, dir_name, new_name)  # 新文件：C://...//数据集//食指弯曲//记录1.1.csv
                    # 提取数据到新的 CSV 文件
                    extract_csv_data(original_df, new_path, i * row_per_part, (i + 1) * row_per_part)
                    print(f"{dir_name}-{new_name}处理完毕")


def pdf_construct_dataset(pdf_path, dataset_path):
    """
    根据整理的手语采样集构建数据集目录，并新建手势数分类的文档。
    执行完后，数据集格式：单手势/动作名，多手势/动作名，并且在pdf_path的同一父目录下新建用手势数分类的文档，其中包含词语名-手势示意图。
    :param pdf_path: "手语采样数据集.pdf"存在的路径
    :param dataset_path: 数据集建立的位置。如：C:/.../数据集
    :return: null
    """
    flag = '多动作'
    doc = fitz.open(pdf_path)
    keys, images = [], []
    simple_key_list, complicate_key_list = [], []
    dictionary = {}
    for page_num, page in enumerate(doc.pages(), start=1):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            # 文字块
            if block["type"] == 0:
                text = block["lines"][0]["spans"][0]["text"]
                font_size = block["lines"][0]["spans"][0]["size"]
                if font_size > 16:  # 建立在词义是二号字、类别是三号字的基础上
                    continue
                else:
                    # 根据flag确定是单动作手语词还是多动作手语词，加入对应的词表中
                    choice = flag in text
                    text = re.split(flag, text)[0]
                    keys.append(text)
                    dictionary[text] = None
                    if choice:
                        complicate_key_list.append(text)
                    else:
                        simple_key_list.append(text)
            # 图片块
            elif block["type"] == 1:
                image = block["image"]
                images.append(BytesIO(image))
        print(f"第{page_num}页提取完成")

    # 建立词-图片对应的字典
    for index, key in enumerate(keys):
        dictionary[key] = images[index]

    docx_path = os.path.join(os.path.dirname(pdf_path), "手语采样数据集-按手势分类.docx")
    construct_dataset(docx_path, dataset_path, dictionary, complicate_key_list, simple_key_list)


def construct_dataset(docx_path, dataset_path, dictionary, complicate_key_list, simple_key_list):
    """
    根据手势名-图片键值对，构建数据集格式、按照手势数分类的采样集.docx
    :param docx_path: docx文档的位置
    :param dataset_path: 建立数据集的位置
    :param dictionary: 字典
    :param complicate_key_list:单动作手势名
    :param simple_key_list: 多动作手势名
    :return:
    """
    # 创建数据集目录结构
    simple_path = os.path.join(dataset_path, "单动作手势")
    complicate_path = os.path.join(dataset_path, "多动作手势")
    os.makedirs(simple_path, exist_ok=True)
    os.makedirs(complicate_path, exist_ok=True)
    # 创建子目录
    for word in simple_key_list:
        obj_path = os.path.join(simple_path, word)
        if os.path.exists(obj_path):
            print(f"'{word}'重复记录")
        else:
            os.makedirs(obj_path)
            print(f"'{word}'成功建立")
    for word in complicate_key_list:
        obj_path = os.path.join(complicate_path, word)
        if os.path.exists(obj_path):
            print(f"'{word}'重复记录")
        else:
            os.makedirs(obj_path)
            print(f"'{word}'成功建立")

    # 新建以手势个数分类的docx
    new_doc = Document()
    # 设置样式
    styles = new_doc.styles
    type_style = styles.add_style("TypeStyle", WD_STYLE_TYPE.PARAGRAPH)  # 大类的字体样式
    type_style.font.name = "Arial"  # 字体
    type_style.font.size = 228600  # 二号字体
    type_style.font.bold = True  # 加粗
    key_style = styles.add_style("KeyStyle", WD_STYLE_TYPE.PARAGRAPH)  # 图片对应词义的字体样式
    key_style.font.name = "Arial"  # 字体
    key_style.font.size = 203200  # 三号字体
    key_style.font.bold = True  # 加粗
    # 写入单手势动作
    new_doc.add_paragraph(text="单手势动作", style="TypeStyle")
    for key in simple_key_list:
        if dictionary[key] is None:
            continue
        new_doc.add_paragraph(key, style="KeyStyle")
        new_doc.add_picture(dictionary[key], width=Cm(14))
        print(f"'{key}'再分类写入成功")
    # 写入多手势动作
    new_doc.add_paragraph(text="多手势动作", style="TypeStyle")
    for key in complicate_key_list:
        if dictionary[key] is None:
            continue
        new_doc.add_paragraph(key, style="KeyStyle")
        new_doc.add_picture(dictionary[key], width=Cm(14))
        print(f"'{key}'再分类写入成功")

    # 保存文档
    new_doc.save(os.path.join(os.path.dirname(docx_path), "手语采样数据集-按手势数分类.docx"))
